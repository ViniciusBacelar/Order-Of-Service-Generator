import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_document(file_name, customer_name, customer_info, computer_info, services):
    doc = docx.Document()
    styles = set_styles()
    add_title(doc, styles)
    add_customer_info(doc, customer_name, customer_info, styles)
    add_computer_info(doc, computer_info, styles)
    add_services(doc, services, styles)
    add_footer(doc, styles)
    doc.save(f'{file_name}.docx')

def set_styles():
    font_name_heading = 'Calibri'
    font_name_body = 'Arial'
    primary_color = RGBColor(0, 123, 255)
    return font_name_heading, font_name_body, primary_color

def add_title(doc, styles):
    font_name_heading, _, _ = styles
    title = doc.add_heading('Ordem de Serviço', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = font_name_heading
    title.runs[0].font.size = Pt(20)
    title.runs[0].bold = True

    company_name = doc.add_paragraph('DB Tecnologia TI')
    company_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_name.runs[0].font.name = font_name_heading
    company_name.runs[0].font.size = Pt(16)

    doc.add_paragraph('')  # Empty line for spacing

def add_customer_info(doc, customer_name, customer_info, styles):
    font_name_heading, _, primary_color = styles
    add_section_heading(doc, 'Dados do Cliente', font_name_heading, primary_color)
    customer_table = doc.add_table(rows=4, cols=2)
    customer_table.style = 'Table Grid'
    customer_table.autofit = True

    cell_data = [
        ('Nome:', customer_name),
        ('Endereço:', customer_info.get('address', '')),
        ('Telefone:', customer_info.get('phone', '')),
        ('Email:', customer_info.get('email', ''))
    ]

    for i, (header, value) in enumerate(cell_data):
        customer_table.cell(i, 0).text = header
        customer_table.cell(i, 1).text = value

def add_computer_info(doc, computer_info, styles):
    font_name_heading, _, primary_color = styles
    doc.add_paragraph('')  # Empty line for spacing
    add_section_heading(doc, 'Informações do Computador', font_name_heading, primary_color)
    computer_info_table = doc.add_table(rows=3, cols=2)
    computer_info_table.style = 'Table Grid'
    computer_info_table.autofit = True

    cell_data = [
        ('Modelo:', computer_info['model']),
        ('Número de Série:', computer_info['serial_number']),
        ('Custo da Manutenção:', f'R$ {computer_info["maintenance_cost"]:}')
    ]

    for i, (header, value) in enumerate(cell_data):
        computer_info_table.cell(i, 0).text = header
        computer_info_table.cell(i, 1).text = value

def add_services(doc, services, styles):
    font_name_heading, _, primary_color = styles
    doc.add_paragraph('')  # Empty line for spacing
    add_section_heading(doc, 'Serviços a serem Realizados', font_name_heading, primary_color)
    services_table = doc.add_table(rows=len(services) + 1, cols=2)
    services_table.style = 'Table Grid'
    services_table.autofit = True
    services_table.cell(0, 0).text = 'Serviço'
    services_table.cell(0, 1).text = 'Descrição'
    for i, (header, value) in enumerate(services):
        services_table.cell(i+1, 0).text = header
        services_table.cell(i+1, 1).text = value

def add_footer(doc, styles):
    _, font_name_body, _ = styles
    doc.add_paragraph('')  # Empty line for spacing
    footer_paragraph = doc.add_paragraph('Obrigado por escolher a DB Tecnologia TI!')
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.runs[0].font.name = font_name_body
    footer_paragraph.runs[0].font.size = Pt(10)

def add_section_heading(doc, text, font_name_heading, primary_color):
    heading = doc.add_heading(text, level=2)
    heading.runs[0].font.name = font_name_heading
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].bold = True
    heading.runs[0].font.color.rgb = primary_color
    return heading

# Example usage:
nome = input("Digite o nome do cliente: ")
endereco = input("Digite o endereço do cliente: ")
telefone = input("Digite o telefone: ")
email = input("Email: ")
modelo = input("Modelo: ")
serial = input("Serial: ")
custo = input("Custo: ")

file_name = 'ordem_de_servico'
customer_name = nome
customer_info = {
    'address': endereco,
    'phone': telefone,
    'email': email
}
computer_info = {
    'model': modelo,
    'serial_number': serial,
    'maintenance_cost': custo
}

# num_services = int(input("Quantos serviços: "))
# services = [(
#     'formatacao','instalar windows 11'
# ),
# ('troca de armazenamento','troca de de HD por um SSD 420gb',
# )]

services=[]
z=input("Quantidade de serviços: ")
for x in range(int(z)):
    name = input(f"Nome do serviço {x+1}: ")
    description = input(f"Descrição do serviço {x+1}: ")
    services.append((name, description))

create_word_document(file_name, customer_name, customer_info, computer_info, services)
